"""
Document Parser Service
Extracts structured text from various document formats
"""

import os
import re
from typing import Dict, List, Optional
from dataclasses import dataclass
from datetime import datetime
import logging

logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Metadata extracted from document"""
    filename: str
    file_type: str
    date: Optional[datetime] = None
    participants: List[str] = None
    topic: Optional[str] = None
    
@dataclass
class TextSegment:
    """A segment of text with optional speaker attribution"""
    content: str
    speaker: Optional[str] = None
    timestamp: Optional[str] = None
    position: int = 0

@dataclass
class ParsedDocument:
    """Parsed document with content and metadata"""
    content: str
    metadata: DocumentMetadata
    segments: List[TextSegment]


class DocumentParserService:
    """
    Service for parsing various document formats into structured text.
    Supports: TXT, MD, DOCX, PDF, EML
    """
    
    def parse_file(self, file_path: str) -> ParsedDocument:
        """Parse a file and return structured document"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        filename = os.path.basename(file_path)
        ext = filename.split('.')[-1].lower()
        
        if ext in ['txt', 'md']:
            return self._parse_text(file_path, filename)
        elif ext == 'docx':
            return self._parse_docx(file_path, filename)
        elif ext == 'pdf':
            return self._parse_pdf(file_path, filename)
        elif ext in ['eml', 'msg']:
            return self._parse_email(file_path, filename)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    
    def _parse_text(self, file_path: str, filename: str) -> ParsedDocument:
        """Parse plain text or markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        metadata = DocumentMetadata(
            filename=filename,
            file_type='text'
        )
        
        # Extract metadata from content
        metadata = self._extract_metadata(content, metadata)
        
        # Segment content
        segments = self._segment_content(content)
        
        return ParsedDocument(
            content=content,
            metadata=metadata,
            segments=segments
        )
    
    def _parse_docx(self, file_path: str, filename: str) -> ParsedDocument:
        """Parse DOCX file"""
        try:
            import docx
            doc = docx.Document(file_path)
            content = '\n'.join([para.text for para in doc.paragraphs])
            
            metadata = DocumentMetadata(
                filename=filename,
                file_type='docx'
            )
            metadata = self._extract_metadata(content, metadata)
            segments = self._segment_content(content)
            
            return ParsedDocument(
                content=content,
                metadata=metadata,
                segments=segments
            )
        except ImportError:
            logger.warning("python-docx not installed, falling back to text parsing")
            return self._parse_text(file_path, filename)
    
    def _parse_pdf(self, file_path: str, filename: str) -> ParsedDocument:
        """Parse PDF file"""
        try:
            import PyPDF2
            content = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
            
            metadata = DocumentMetadata(
                filename=filename,
                file_type='pdf'
            )
            metadata = self._extract_metadata(content, metadata)
            segments = self._segment_content(content)
            
            return ParsedDocument(
                content=content,
                metadata=metadata,
                segments=segments
            )
        except ImportError:
            logger.warning("PyPDF2 not installed, falling back to text parsing")
            return self._parse_text(file_path, filename)
    
    def _parse_email(self, file_path: str, filename: str) -> ParsedDocument:
        """Parse email file"""
        try:
            from email import message_from_file
            with open(file_path, 'r') as f:
                msg = message_from_file(f)
            
            # Extract email content
            if msg.is_multipart():
                content = ""
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        content += part.get_payload(decode=True).decode('utf-8', errors='ignore')
            else:
                content = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            
            metadata = DocumentMetadata(
                filename=filename,
                file_type='email',
                date=self._parse_email_date(msg.get('Date')),
                participants=[msg.get('From'), msg.get('To')],
                topic=msg.get('Subject')
            )
            
            segments = self._segment_content(content)
            
            return ParsedDocument(
                content=content,
                metadata=metadata,
                segments=segments
            )
        except Exception as e:
            logger.warning(f"Email parsing failed: {e}, falling back to text parsing")
            return self._parse_text(file_path, filename)
    
    def _extract_metadata(self, content: str, metadata: DocumentMetadata) -> DocumentMetadata:
        """Extract metadata from content"""
        # Extract date
        date_pattern = r'\b(\d{4}-\d{2}-\d{2}|\d{2}/\d{2}/\d{4})\b'
        date_match = re.search(date_pattern, content)
        if date_match:
            try:
                date_str = date_match.group(1)
                if '-' in date_str:
                    metadata.date = datetime.strptime(date_str, '%Y-%m-%d')
                else:
                    metadata.date = datetime.strptime(date_str, '%m/%d/%Y')
            except:
                pass
        
        # Extract participants (names followed by colons, common in meeting notes)
        participant_pattern = r'^([A-Z][a-z]+ [A-Z][a-z]+):'
        participants = set()
        for line in content.split('\n'):
            match = re.match(participant_pattern, line.strip())
            if match:
                participants.add(match.group(1))
        
        if participants:
            metadata.participants = list(participants)
        
        # Extract topic from first line or title
        lines = [l.strip() for l in content.split('\n') if l.strip()]
        if lines:
            first_line = lines[0]
            if len(first_line) < 100 and not first_line.endswith('.'):
                metadata.topic = first_line
        
        return metadata
    
    def _segment_content(self, content: str) -> List[TextSegment]:
        """Segment content by speaker or paragraph"""
        segments = []
        
        # Try to detect speaker-based format (Name: text)
        speaker_pattern = r'^([A-Z][a-z]+ [A-Z][a-z]+):\s*(.+)$'
        
        position = 0
        for line in content.split('\n'):
            line = line.strip()
            if not line:
                continue
            
            match = re.match(speaker_pattern, line)
            if match:
                segments.append(TextSegment(
                    content=match.group(2),
                    speaker=match.group(1),
                    position=position
                ))
            else:
                segments.append(TextSegment(
                    content=line,
                    position=position
                ))
            
            position += 1
        
        # If no speaker-based segments, split by paragraphs
        if not any(s.speaker for s in segments):
            segments = []
            paragraphs = content.split('\n\n')
            for i, para in enumerate(paragraphs):
                if para.strip():
                    segments.append(TextSegment(
                        content=para.strip(),
                        position=i
                    ))
        
        return segments
    
    def _parse_email_date(self, date_str: Optional[str]) -> Optional[datetime]:
        """Parse email date string"""
        if not date_str:
            return None
        
        try:
            from email.utils import parsedate_to_datetime
            return parsedate_to_datetime(date_str)
        except:
            return None
    
    def _get_context(self, text: str, position: tuple, window: int = 100) -> str:
        """Get context around a position in text"""
        start = max(0, position[0] - window)
        end = min(len(text), position[1] + window)
        return text[start:end]
